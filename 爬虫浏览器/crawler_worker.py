import os
import json
import re
import logging
from datetime import datetime
from urllib.parse import urlparse, urljoin
import requests
from bs4 import BeautifulSoup
from utils import SELENIUM_AVAILABLE, DOCX_AVAILABLE

if SELENIUM_AVAILABLE:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.common.by import By
    from webdriver_manager.chrome import ChromeDriverManager
    from selenium.webdriver.chrome.service import Service

if DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

class CrawlerWorker:
    """增强版爬虫引擎，支持动态渲染和静态解析"""
    
    def __init__(self, output_dir="crawled_data"):
        self.output_dir = output_dir
        self.crawled_data = []
        self.driver = None
        self.setup_logging()
        if SELENIUM_AVAILABLE:
            self.setup_selenium()

    def setup_logging(self):
        os.makedirs(self.output_dir, exist_ok=True)
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(os.path.join(self.output_dir, 'crawler.log')),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def setup_selenium(self):
        try:
            chrome_options = Options()
            chrome_options.add_argument("--headless")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument(
                "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            )

            service = Service(ChromeDriverManager().install())
            self.driver = webdriver.Chrome(service=service, options=chrome_options)
            self.logger.info("Selenium 初始化成功")
        except Exception as e:
            self.logger.warning(f"Selenium 初始化失败: {e}")
            self.driver = None

    def is_valid_url(self, url):
        try:
            parsed = urlparse(url)
            return bool(parsed.netloc and parsed.scheme in ['http', 'https'])
        except:
            return False

    def can_fetch(self, url, user_agent="*"):
        try:
            parsed = urlparse(url)
            base_url = f"{parsed.scheme}://{parsed.netloc}"
            robots_url = f"{base_url}/robots.txt"

            rp = urllib.robotparser.RobotFileParser()
            rp.set_url(robots_url)
            rp.read()
            return rp.can_fetch(user_agent, url)
        except Exception as e:
            self.logger.warning(f"无法检查 robots.txt: {e}")
            return True  # 宽松策略

    def clean_text(self, text):
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'http\S+', '', text)  # 移除 URL
        return text.strip()

    def extract_page_data(self, soup, current_url):
        # 移除无关标签
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        title = soup.find('title')
        title_text = title.get_text().strip() if title else "无标题"

        # 提取正文
        paragraphs = []
        for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
            txt = p.get_text().strip()
            if len(txt) > 20:
                paragraphs.append(txt)

        full_text = '\n'.join(paragraphs)
        clean_text = self.clean_text(full_text)

        links = [urljoin(current_url, a.get('href')) for a in soup.find_all('a', href=True)]
        internal_links = [link for link in links if urlparse(link).netloc == urlparse(current_url).netloc]
        external_links = [link for link in links if link not in internal_links]

        return {
            "title": title_text,
            "url": current_url,
            "timestamp": datetime.now().isoformat(),
            "content_preview": clean_text[:1000] + "..." if len(clean_text) > 1000 else clean_text,
            "full_content": clean_text,
            "word_count": len(clean_text.split()),
            "char_count": len(clean_text),
            "total_links": len(links),
            "internal_links": len(internal_links),
            "external_links": len(external_links),
            "top_links": links[:50],
            "images": [img.get('src') for img in soup.find_all('img', src=True)][:20],
            "meta_description": "",
        }

    def crawl_with_requests(self, url):
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            if response.encoding != 'utf-8':
                response.encoding = 'utf-8'
            soup = BeautifulSoup(response.content, 'html.parser')
            return soup, True
        except Exception as e:
            self.logger.warning(f"Requests 失败: {e}")
            return None, False

    def crawl_with_selenium(self, url):
        if not self.driver:
            return None, False
        try:
            self.driver.get(url)
            WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            page_source = self.driver.page_source
            soup = BeautifulSoup(page_source, 'html.parser')
            return soup, True
        except Exception as e:
            self.logger.warning(f"Selenium 失败: {e}")
            return None, False

    def crawl_single_page(self, url):
        if not self.is_valid_url(url):
            return False, "无效的URL"

        success, msg = False, "未知错误"
        soup, success = self.crawl_with_requests(url)
        if not success and SELENIUM_AVAILABLE:
            soup, success = self.crawl_with_selenium(url)

        if not success or not soup:
            return False, "页面获取失败"

        data = self.extract_page_data(soup, url)
        self.crawled_data.append(data)
        self.save_data()
        return True, f"成功抓取 {len(data['full_content'])} 字符"

    def save_data(self):
        try:
            # JSON 全量存储
            with open(os.path.join(self.output_dir, 'crawled_data.json'), 'w', encoding='utf-8') as f:
                json.dump(self.crawled_data, f, ensure_ascii=False, indent=2)

            # TXT 训练数据格式
            with open(os.path.join(self.output_dir, 'training_data.txt'), 'w', encoding='utf-8') as f:
                for item in self.crawled_data:
                    f.write(f"URL: {item['url']}\n")
                    f.write(f"标题: {item['title']}\n")
                    f.write(f"字数: {item['word_count']} 字\n")
                    f.write(f"内容:\n{item['full_content']}\n")
                    f.write("\n" + "=" * 80 + "\n\n")
        except Exception as e:
            self.logger.error(f"保存失败: {e}")

    def export_to_docx(self, filepath):
        if not DOCX_AVAILABLE:
            return False, "DOCX库未安装"
        if not self.crawled_data:
            return False, "无数据可导出"

        try:
            doc = Document()
            doc.add_heading('AI数据采集报告', 0)

            for i, d in enumerate(self.crawled_data, 1):
                doc.add_heading(f"{i}. {d['title']}", level=1)
                p = doc.add_paragraph("")
                p.add_run(f"来源: ").bold = True
                p.add_run(d['url'])
                doc.add_paragraph(f"采集时间: {d['timestamp']}")
                doc.add_heading("内容摘要", level=2)
                doc.add_paragraph(d['full_content'])
                doc.add_page_break()

            doc.save(filepath)
            return True, f"已导出至 {filepath}"
        except Exception as e:
            return False, str(e)