import sys
import os
import json
import urllib.robotparser
from urllib.parse import urlparse, urljoin
from datetime import datetime
import time
import re
import logging
from threading import Thread
from PyQt5.QtCore import QUrl, Qt, QTimer, pyqtSignal, QSize, QStandardPaths
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QVBoxLayout, QHBoxLayout,
    QWidget, QPushButton, QLineEdit, QToolBar, QLabel,
    QTabWidget, QStatusBar, QAction, QFileDialog, QMessageBox,
    QSizePolicy, QListWidget, QListWidgetItem, QTextEdit, QSplitter,
    QGroupBox, QProgressBar, QMenu, QDialog, QDialogButtonBox,
    QListWidget, QTreeWidget, QTreeWidgetItem, QHeaderView, QCheckBox
)
from PyQt5.QtGui import QFont, QIcon, QPixmap, QPalette, QColor
from PyQt5.QtWebEngineWidgets import QWebEngineView, QWebEnginePage, QWebEngineProfile, QWebEngineDownloadItem, QWebEngineSettings
import requests
from bs4 import BeautifulSoup

# 尝试导入高级依赖（非强制）
try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.common.by import By
    from webdriver_manager.chrome import ChromeDriverManager
    from selenium.webdriver.chrome.service import Service
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class CustomWebEnginePage(QWebEnginePage):
    """自定义网页引擎页面，处理导航请求和新窗口请求"""
    
    def __init__(self, parent=None, main_window=None):
        super().__init__(parent)
        self.main_window = main_window
        self.navigation_attempts = {}

    def acceptNavigationRequest(self, url, navigation_type, isMainFrame):
        """处理导航请求，允许所有类型的导航"""
        # 记录导航尝试
        url_str = url.toString()
        self.navigation_attempts[url_str] = self.navigation_attempts.get(url_str, 0) + 1
        
        # 允许所有导航请求
        print(f"导航请求: {url_str}, 类型: {navigation_type}, 主框架: {isMainFrame}")
        return True

    def createWindow(self, type):
        """创建新窗口/新标签页 - 这是关键函数，处理新窗口请求"""
        print(f"创建新窗口请求: {type}")
        
        if self.main_window:
            # 在主窗口中创建新标签页
            new_browser = self.main_window.add_new_tab(QUrl("about:blank"), "新标签页")
            return new_browser.page()
        
        # 如果没有主窗口引用，创建一个新的浏览器窗口
        new_browser = QWebEngineView()
        new_page = CustomWebEnginePage(new_browser)
        new_browser.setPage(new_page)
        return new_page


class DownloadManager(QDialog):
    """下载管理器"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("下载管理器")
        self.setGeometry(300, 300, 800, 500)
        self.setup_ui()
        self.downloads = []

    def setup_ui(self):
        layout = QVBoxLayout(self)

        # 工具栏
        toolbar = QHBoxLayout()
        self.clear_btn = QPushButton("清空已完成")
        self.open_folder_btn = QPushButton("打开下载文件夹")
        self.pause_all_btn = QPushButton("暂停全部")
        self.resume_all_btn = QPushButton("继续全部")

        self.clear_btn.clicked.connect(self.clear_completed)
        self.open_folder_btn.clicked.connect(self.open_download_folder)
        self.pause_all_btn.clicked.connect(self.pause_all)
        self.resume_all_btn.clicked.connect(self.resume_all)

        toolbar.addWidget(self.clear_btn)
        toolbar.addWidget(self.open_folder_btn)
        toolbar.addWidget(self.pause_all_btn)
        toolbar.addWidget(self.resume_all_btn)
        toolbar.addStretch()

        # 下载列表
        self.download_list = QTreeWidget()
        self.download_list.setHeaderLabels(["文件名", "进度", "状态", "大小", "速度", "剩余时间"])
        self.download_list.header().setSectionResizeMode(0, QHeaderView.Stretch)

        layout.addLayout(toolbar)
        layout.addWidget(self.download_list)

    def add_download(self, download_item):
        """添加新的下载项"""
        item = QTreeWidgetItem(self.download_list)
        item.download = download_item
        self.downloads.append(item)
        
        filename = os.path.basename(download_item.path())
        item.setText(0, filename)
        item.setText(1, "0%")
        item.setText(2, "下载中")
        item.setText(3, "未知")
        item.setText(4, "0 KB/s")
        item.setText(5, "未知")
        
        # 连接信号
        download_item.downloadProgress.connect(lambda bytes_received, bytes_total: 
                                             self.update_progress(item, bytes_received, bytes_total))
        download_item.finished.connect(lambda: self.download_finished(item))
        
        self.download_list.addTopLevelItem(item)

    def update_progress(self, item, bytes_received, bytes_total):
        """更新下载进度"""
        if bytes_total > 0:
            percent = int((bytes_received / bytes_total) * 100)
            item.setText(1, f"{percent}%")
            
            # 计算下载速度（简化版）
            speed = "计算中..."
            time_left = "计算中..."
            
            item.setText(2, "下载中")
            item.setText(3, f"{bytes_received//1024}KB / {bytes_total//1024}KB")
            item.setText(4, speed)
            item.setText(5, time_left)

    def download_finished(self, item):
        """下载完成"""
        if item.download.state() == QWebEngineDownloadItem.DownloadCompleted:
            item.setText(1, "100%")
            item.setText(2, "已完成")
            item.setText(4, "")
            item.setText(5, "")
        else:
            item.setText(2, "失败")

    def clear_completed(self):
        """清除已完成的下载"""
        for i in range(self.download_list.topLevelItemCount() - 1, -1, -1):
            item = self.download_list.topLevelItem(i)
            if item.text(2) in ["已完成", "失败"]:
                self.download_list.takeTopLevelItem(i)

    def open_download_folder(self):
        """打开下载文件夹"""
        download_path = QStandardPaths.writableLocation(QStandardPaths.DownloadLocation)
        if os.path.exists(download_path):
            if os.name == 'nt':  # Windows
                os.startfile(download_path)
            elif os.name == 'posix':  # Linux or macOS
                if sys.platform == 'darwin':  # macOS
                    os.system(f'open "{download_path}"')
                else:  # Linux
                    os.system(f'xdg-open "{download_path}"')

    def pause_all(self):
        """暂停所有下载"""
        for item in self.downloads:
            if hasattr(item, 'download') and item.download.state() == QWebEngineDownloadItem.DownloadInProgress:
                item.download.pause()
                item.setText(2, "已暂停")

    def resume_all(self):
        """继续所有下载"""
        for item in self.downloads:
            if hasattr(item, 'download') and item.download.state() == QWebEngineDownloadItem.DownloadPaused:
                item.download.resume()
                item.setText(2, "下载中")


class HistoryManager(QDialog):
    """历史记录管理器"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("浏览历史")
        self.setGeometry(300, 300, 800, 500)
        self.setup_ui()
        self.history = []

    def setup_ui(self):
        layout = QVBoxLayout(self)

        # 工具栏
        toolbar = QHBoxLayout()
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("搜索历史记录...")
        self.search_box.textChanged.connect(self.filter_history)
        
        self.clear_btn = QPushButton("清空历史")
        self.clear_btn.clicked.connect(self.clear_history)

        toolbar.addWidget(QLabel("搜索:"))
        toolbar.addWidget(self.search_box)
        toolbar.addStretch()
        toolbar.addWidget(self.clear_btn)

        # 历史列表
        self.history_list = QTreeWidget()
        self.history_list.setHeaderLabels(["标题", "网址", "访问时间"])
        self.history_list.header().setSectionResizeMode(0, QHeaderView.Stretch)
        self.history_list.itemDoubleClicked.connect(self.open_history_item)

        layout.addLayout(toolbar)
        layout.addWidget(self.history_list)

    def add_history(self, title, url):
        """添加历史记录"""
        self.history.append({
            "title": title,
            "url": url,
            "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        self.refresh_list()

    def filter_history(self, text):
        """过滤历史记录"""
        for i in range(self.history_list.topLevelItemCount()):
            item = self.history_list.topLevelItem(i)
            match = (text.lower() in item.text(0).lower() or 
                    text.lower() in item.text(1).lower())
            item.setHidden(not match)

    def refresh_list(self):
        """刷新列表"""
        self.history_list.clear()
        for record in self.history:
            item = QTreeWidgetItem(self.history_list)
            item.setText(0, record["title"])
            item.setText(1, record["url"])
            item.setText(2, record["time"])
            self.history_list.addTopLevelItem(item)

    def open_history_item(self, item, column):
        """打开历史记录项"""
        url = item.text(1)
        if self.parent():
            self.parent().add_new_tab(QUrl(url))

    def clear_history(self):
        """清空历史记录"""
        if QMessageBox.question(self, "确认", "清空所有历史记录？") == QMessageBox.Yes:
            self.history.clear()
            self.history_list.clear()


class BookmarksManager(QDialog):
    """书签管理器"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("书签管理器")
        self.setGeometry(300, 300, 800, 500)
        self.setup_ui()
        self.bookmarks = []

    def setup_ui(self):
        layout = QVBoxLayout(self)

        # 工具栏
        toolbar = QHBoxLayout()
        self.add_btn = QPushButton("添加书签")
        self.delete_btn = QPushButton("删除书签")
        self.import_btn = QPushButton("导入")
        self.export_btn = QPushButton("导出")

        self.add_btn.clicked.connect(self.add_bookmark)
        self.delete_btn.clicked.connect(self.delete_bookmark)
        self.import_btn.clicked.connect(self.import_bookmarks)
        self.export_btn.clicked.connect(self.export_bookmarks)

        toolbar.addWidget(self.add_btn)
        toolbar.addWidget(self.delete_btn)
        toolbar.addWidget(self.import_btn)
        toolbar.addWidget(self.export_btn)
        toolbar.addStretch()

        # 书签列表
        self.bookmarks_list = QTreeWidget()
        self.bookmarks_list.setHeaderLabels(["标题", "网址", "添加时间"])
        self.bookmarks_list.header().setSectionResizeMode(0, QHeaderView.Stretch)
        self.bookmarks_list.itemDoubleClicked.connect(self.open_bookmark)

        layout.addLayout(toolbar)
        layout.addWidget(self.bookmarks_list)

    def add_bookmark(self, title="", url=""):
        """添加书签"""
        if not title or not url:
            # 从父窗口获取当前页面信息
            if self.parent():
                browser = self.parent().tab_widget.currentWidget()
                if isinstance(browser, QWebEngineView):
                    title = browser.title()
                    url = browser.url().toString()
        
        if title and url:
            self.bookmarks.append({
                "title": title,
                "url": url,
                "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            })
            self.refresh_list()

    def delete_bookmark(self):
        """删除书签"""
        current_item = self.bookmarks_list.currentItem()
        if current_item:
            index = self.bookmarks_list.indexOfTopLevelItem(current_item)
            if index >= 0:
                self.bookmarks.pop(index)
                self.refresh_list()

    def open_bookmark(self, item, column):
        """打开书签"""
        url = item.text(1)
        if self.parent():
            self.parent().add_new_tab(QUrl(url))

    def import_bookmarks(self):
        """导入书签"""
        path, _ = QFileDialog.getOpenFileName(self, "导入书签", "", "JSON文件 (*.json)")
        if path:
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    self.bookmarks = json.load(f)
                self.refresh_list()
                QMessageBox.information(self, "成功", "书签导入成功")
            except Exception as e:
                QMessageBox.critical(self, "失败", f"导入失败: {e}")

    def export_bookmarks(self):
        """导出书签"""
        path, _ = QFileDialog.getSaveFileName(self, "导出书签", "", "JSON文件 (*.json)")
        if path:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(self.bookmarks, f, ensure_ascii=False, indent=2)
                QMessageBox.information(self, "成功", "书签导出成功")
            except Exception as e:
                QMessageBox.critical(self, "失败", f"导出失败: {e}")

    def refresh_list(self):
        """刷新列表"""
        self.bookmarks_list.clear()
        for bookmark in self.bookmarks:
            item = QTreeWidgetItem(self.bookmarks_list)
            item.setText(0, bookmark["title"])
            item.setText(1, bookmark["url"])
            item.setText(2, bookmark["time"])
            self.bookmarks_list.addTopLevelItem(item)


class SettingsDialog(QDialog):
    """设置对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("浏览器设置")
        self.setGeometry(400, 400, 600, 400)
        self.setup_ui()

    def setup_ui(self):
        layout = QVBoxLayout(self)

        # 下载设置
        download_group = QGroupBox("下载设置")
        download_layout = QVBoxLayout()
        
        self.download_path_edit = QLineEdit()
        self.download_path_edit.setText(QStandardPaths.writableLocation(QStandardPaths.DownloadLocation))
        self.browse_btn = QPushButton("浏览...")
        self.browse_btn.clicked.connect(self.browse_download_path)
        
        path_layout = QHBoxLayout()
        path_layout.addWidget(QLabel("下载路径:"))
        path_layout.addWidget(self.download_path_edit)
        path_layout.addWidget(self.browse_btn)
        
        self.ask_save_check = QCheckBox("每次下载前询问保存位置")
        self.ask_save_check.setChecked(True)
        
        download_layout.addLayout(path_layout)
        download_layout.addWidget(self.ask_save_check)
        download_group.setLayout(download_layout)

        # 隐私设置
        privacy_group = QGroupBox("隐私设置")
        privacy_layout = QVBoxLayout()
        
        self.clear_on_exit = QCheckBox("退出时清除浏览数据")
        self.block_images = QCheckBox("阻止图片加载（加速浏览）")
        self.javascript_enabled = QCheckBox("启用JavaScript")
        self.javascript_enabled.setChecked(True)
        
        privacy_layout.addWidget(self.clear_on_exit)
        privacy_layout.addWidget(self.block_images)
        privacy_layout.addWidget(self.javascript_enabled)
        privacy_group.setLayout(privacy_layout)

        # 按钮
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)

        layout.addWidget(download_group)
        layout.addWidget(privacy_group)
        layout.addStretch()
        layout.addWidget(button_box)

    def browse_download_path(self):
        """选择下载路径"""
        path = QFileDialog.getExistingDirectory(self, "选择下载文件夹", self.download_path_edit.text())
        if path:
            self.download_path_edit.setText(path)


class CrawlerWorker:
    """增强版爬虫引擎，支持动态渲染和静态解析"""
    
    def __init__(self, output_dir="crawled_data"):
        self.output_dir = output_dir
        self.crawled_data = []
        self.driver = None
        self.setup_logging()
        if SELENIUM_AVAILABLE:
            self.setup_selenium()

    def setup_logging(self):
        os.makedirs(self.output_dir, exist_ok=True)
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(os.path.join(self.output_dir, 'crawler.log')),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def setup_selenium(self):
        try:
            chrome_options = Options()
            chrome_options.add_argument("--headless")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--window-size=1920,1080")
            chrome_options.add_argument(
                "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            )

            service = Service(ChromeDriverManager().install())
            self.driver = webdriver.Chrome(service=service, options=chrome_options)
            self.logger.info("Selenium 初始化成功")
        except Exception as e:
            self.logger.warning(f"Selenium 初始化失败: {e}")
            self.driver = None

    def is_valid_url(self, url):
        try:
            parsed = urlparse(url)
            return bool(parsed.netloc and parsed.scheme in ['http', 'https'])
        except:
            return False

    def can_fetch(self, url, user_agent="*"):
        try:
            parsed = urlparse(url)
            base_url = f"{parsed.scheme}://{parsed.netloc}"
            robots_url = f"{base_url}/robots.txt"

            rp = urllib.robotparser.RobotFileParser()
            rp.set_url(robots_url)
            rp.read()
            return rp.can_fetch(user_agent, url)
        except Exception as e:
            self.logger.warning(f"无法检查 robots.txt: {e}")
            return True  # 宽松策略

    def clean_text(self, text):
        if not text:
            return ""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'http\S+', '', text)  # 移除 URL
        return text.strip()

    def extract_page_data(self, soup, current_url):
        # 移除无关标签
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        title = soup.find('title')
        title_text = title.get_text().strip() if title else "无标题"

        # 提取正文
        paragraphs = []
        for p in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
            txt = p.get_text().strip()
            if len(txt) > 20:
                paragraphs.append(txt)

        full_text = '\n'.join(paragraphs)
        clean_text = self.clean_text(full_text)

        links = [urljoin(current_url, a.get('href')) for a in soup.find_all('a', href=True)]
        internal_links = [link for link in links if urlparse(link).netloc == urlparse(current_url).netloc]
        external_links = [link for link in links if link not in internal_links]

        return {
            "title": title_text,
            "url": current_url,
            "timestamp": datetime.now().isoformat(),
            "content_preview": clean_text[:1000] + "..." if len(clean_text) > 1000 else clean_text,
            "full_content": clean_text,
            "word_count": len(clean_text.split()),
            "char_count": len(clean_text),
            "total_links": len(links),
            "internal_links": len(internal_links),
            "external_links": len(external_links),
            "top_links": links[:50],
            "images": [img.get('src') for img in soup.find_all('img', src=True)][:20],
            "meta_description": "",
        }

    def crawl_with_requests(self, url):
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            if response.encoding != 'utf-8':
                response.encoding = 'utf-8'
            soup = BeautifulSoup(response.content, 'html.parser')
            return soup, True
        except Exception as e:
            self.logger.warning(f"Requests 失败: {e}")
            return None, False

    def crawl_with_selenium(self, url):
        if not self.driver:
            return None, False
        try:
            self.driver.get(url)
            WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            page_source = self.driver.page_source
            soup = BeautifulSoup(page_source, 'html.parser')
            return soup, True
        except Exception as e:
            self.logger.warning(f"Selenium 失败: {e}")
            return None, False

    def crawl_single_page(self, url):
        if not self.is_valid_url(url):
            return False, "无效的URL"

        success, msg = False, "未知错误"
        soup, success = self.crawl_with_requests(url)
        if not success and SELENIUM_AVAILABLE:
            soup, success = self.crawl_with_selenium(url)

        if not success or not soup:
            return False, "页面获取失败"

        data = self.extract_page_data(soup, url)
        self.crawled_data.append(data)
        self.save_data()
        return True, f"成功抓取 {len(data['full_content'])} 字符"

    def save_data(self):
        try:
            # JSON 全量存储
            with open(os.path.join(self.output_dir, 'crawled_data.json'), 'w', encoding='utf-8') as f:
                json.dump(self.crawled_data, f, ensure_ascii=False, indent=2)

            # TXT 训练数据格式
            with open(os.path.join(self.output_dir, 'training_data.txt'), 'w', encoding='utf-8') as f:
                for item in self.crawled_data:
                    f.write(f"URL: {item['url']}\n")
                    f.write(f"标题: {item['title']}\n")
                    f.write(f"字数: {item['word_count']} 字\n")
                    f.write(f"内容:\n{item['full_content']}\n")
                    f.write("\n" + "=" * 80 + "\n\n")
        except Exception as e:
            self.logger.error(f"保存失败: {e}")

    def export_to_docx(self, filepath):
        if not DOCX_AVAILABLE:
            return False, "DOCX库未安装"
        if not self.crawled_data:
            return False, "无数据可导出"

        try:
            doc = Document()
            doc.add_heading('AI数据采集报告', 0)

            for i, d in enumerate(self.crawled_data, 1):
                doc.add_heading(f"{i}. {d['title']}", level=1)
                p = doc.add_paragraph("")
                p.add_run(f"来源: ").bold = True
                p.add_run(d['url'])
                doc.add_paragraph(f"采集时间: {d['timestamp']}")
                doc.add_heading("内容摘要", level=2)
                doc.add_paragraph(d['full_content'])
                doc.add_page_break()

            doc.save(filepath)
            return True, f"已导出至 {filepath}"
        except Exception as e:
            return False, str(e)


class ModernBrowser(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("道衍AI浏览器 - 智能合规爬虫版")
        self.resize(1400, 900)
        
        # 初始化组件
        self.crawler = CrawlerWorker()
        self.download_manager = DownloadManager(self)
        self.history_manager = HistoryManager(self)
        self.bookmarks_manager = BookmarksManager(self)
        self.settings_dialog = SettingsDialog(self)
        
        # 下载设置
        self.download_path = QStandardPaths.writableLocation(QStandardPaths.DownloadLocation)
        self.ask_before_download = True

        self.setup_ui()
        self.setup_downloads()

    def setup_downloads(self):
        """设置下载处理器"""
        profile = QWebEngineProfile.defaultProfile()
        profile.downloadRequested.connect(self.on_download_requested)

    def on_download_requested(self, download):
        """处理下载请求"""
        if self.ask_before_download:
            path, _ = QFileDialog.getSaveFileName(
                self, "保存文件", 
                os.path.join(self.download_path, download.downloadFileName())
            )
            if path:
                download.setPath(path)
                download.accept()
                self.download_manager.add_download(download)
                self.download_manager.show()
            else:
                download.cancel()
        else:
            download.setPath(os.path.join(self.download_path, download.downloadFileName()))
            download.accept()
            self.download_manager.add_download(download)
            self.download_manager.show()

    def setup_ui(self):
        self.setStyleSheet("""
            QMainWindow { background-color: #f0f0f0; }
            QLineEdit {
                padding: 10px;
                border: 2px solid #ccc;
                border-radius: 20px;
                font-size: 14px;
            }
            QPushButton {
                background-color: white;
                border: 1px solid #ddd;
                padding: 8px 12px;
                border-radius: 18px;
                font-size: 13px;
            }
            QPushButton:hover { background-color: #f9f9f9; border-color: #aaa; }
            QTabBar::tab {
                background-color: white;
                color: #333;
                padding: 8px 16px;
                margin-right: 2px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QTabBar::tab:selected {
                background-color: #e6f7ff;
                font-weight: bold;
            }
            QToolBar {
                background-color: white;
                border-bottom: 1px solid #ddd;
                spacing: 10px;
                padding: 8px;
            }
            QLabel#status {
                color: #555;
                padding: 4px 8px;
                background: #eee;
                border-radius: 10px;
            }
            QListWidget {
                border: 1px solid #ddd;
                border-radius: 8px;
                background: white;
            }
            QTextEdit {
                border: 1px solid #ddd;
                border-radius: 8px;
                background: white;
            }
            QGroupBox {
                border: 2px solid #ddd;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
                font-weight: bold;
                background: white;
            }
        """)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # 左侧：浏览器主体
        browser_container = QWidget()
        browser_layout = QVBoxLayout(browser_container)
        browser_layout.setContentsMargins(0, 0, 0, 0)
        self.create_toolbar(browser_layout)
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)
        self.tab_widget.currentChanged.connect(self.on_tab_changed)
        browser_layout.addWidget(self.tab_widget)
        self.add_new_tab(QUrl("https://www.baidu.com"), "首页")

        # 右侧：爬虫数据面板
        self.data_panel = self.create_data_panel()

        # 分割器
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(browser_container)
        splitter.addWidget(self.data_panel)
        splitter.setSizes([900, 500])
        main_layout.addWidget(splitter)

        # 状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_label = QLabel("准备就绪", objectName="status")
        self.status_bar.addPermanentWidget(self.status_label)

        # 创建菜单栏
        self.create_menubar()

    def create_menubar(self):
        """创建菜单栏"""
        menubar = self.menuBar()

        # 文件菜单
        file_menu = menubar.addMenu("文件")
        
        new_tab_action = QAction("新建标签页", self)
        new_tab_action.setShortcut("Ctrl+T")
        new_tab_action.triggered.connect(lambda: self.add_new_tab(QUrl("https://www.baidu.com")))
        
        new_window_action = QAction("新建窗口", self)
        new_window_action.setShortcut("Ctrl+N")
        new_window_action.triggered.connect(self.new_window)
        
        quit_action = QAction("退出", self)
        quit_action.setShortcut("Ctrl+Q")
        quit_action.triggered.connect(self.close)
        
        file_menu.addAction(new_tab_action)
        file_menu.addAction(new_window_action)
        file_menu.addSeparator()
        file_menu.addAction(quit_action)

        # 编辑菜单
        edit_menu = menubar.addMenu("编辑")
        
        cut_action = QAction("剪切", self)
        cut_action.setShortcut("Ctrl+X")
        cut_action.triggered.connect(self.cut)
        
        copy_action = QAction("复制", self)
        copy_action.setShortcut("Ctrl+C")
        copy_action.triggered.connect(self.copy)
        
        paste_action = QAction("粘贴", self)
        paste_action.setShortcut("Ctrl+V")
        paste_action.triggered.connect(self.paste)
        
        edit_menu.addAction(cut_action)
        edit_menu.addAction(copy_action)
        edit_menu.addAction(paste_action)

        # 查看菜单
        view_menu = menubar.addMenu("查看")
        
        zoom_in_action = QAction("放大", self)
        zoom_in_action.setShortcut("Ctrl++")
        zoom_in_action.triggered.connect(self.zoom_in)
        
        zoom_out_action = QAction("缩小", self)
        zoom_out_action.setShortcut("Ctrl+-")
        zoom_out_action.triggered.connect(self.zoom_out)
        
        zoom_reset_action = QAction("重置缩放", self)
        zoom_reset_action.setShortcut("Ctrl+0")
        zoom_reset_action.triggered.connect(self.zoom_reset)
        
        view_menu.addAction(zoom_in_action)
        view_menu.addAction(zoom_out_action)
        view_menu.addAction(zoom_reset_action)

        # 书签菜单
        bookmarks_menu = menubar.addMenu("书签")
        
        add_bookmark_action = QAction("添加书签", self)
        add_bookmark_action.setShortcut("Ctrl+D")
        add_bookmark_action.triggered.connect(self.add_bookmark)
        
        bookmarks_manager_action = QAction("书签管理器", self)
        bookmarks_manager_action.setShortcut("Ctrl+Shift+O")
        bookmarks_manager_action.triggered.connect(self.bookmarks_manager.show)
        
        bookmarks_menu.addAction(add_bookmark_action)
        bookmarks_menu.addAction(bookmarks_manager_action)

        # 工具菜单
        tools_menu = menubar.addMenu("工具")
        
        downloads_action = QAction("下载管理器", self)
        downloads_action.setShortcut("Ctrl+J")
        downloads_action.triggered.connect(self.download_manager.show)
        
        history_action = QAction("历史记录", self)
        history_action.setShortcut("Ctrl+H")
        history_action.triggered.connect(self.history_manager.show)
        
        settings_action = QAction("设置", self)
        settings_action.setShortcut("Ctrl+,")
        settings_action.triggered.connect(self.settings_dialog.exec_)
        
        tools_menu.addAction(downloads_action)
        tools_menu.addAction(history_action)
        tools_menu.addSeparator()
        tools_menu.addAction(settings_action)

        # 帮助菜单
        help_menu = menubar.addMenu("帮助")
        
        about_action = QAction("关于", self)
        about_action.triggered.connect(self.show_about)
        
        help_menu.addAction(about_action)

    def create_toolbar(self, parent_layout):
        toolbar = QToolBar()
        toolbar.setMovable(False)

        self.back_btn = QPushButton("⏪")
        self.forward_btn = QPushButton("⏩")
        self.reload_btn = QPushButton("🔄")
        self.home_btn = QPushButton("🏠")
        self.downloads_btn = QPushButton("📥")
        self.history_btn = QPushButton("📚")
        self.bookmarks_btn = QPushButton("🔖")
        
        self.url_bar = QLineEdit()
        self.url_bar.setPlaceholderText("输入网址或关键词（自动百度搜索）")
        
        self.crawl_btn = QPushButton("🕷️ 抓取当前页")
        new_tab_btn = QPushButton("+")

        # 连接信号
        self.back_btn.clicked.connect(self.on_back_clicked)
        self.forward_btn.clicked.connect(self.on_forward_clicked)
        self.reload_btn.clicked.connect(self.on_reload_clicked)
        self.home_btn.clicked.connect(self.go_home)
        self.downloads_btn.clicked.connect(self.download_manager.show)
        self.history_btn.clicked.connect(self.history_manager.show)
        self.bookmarks_btn.clicked.connect(self.bookmarks_manager.show)
        self.url_bar.returnPressed.connect(self.on_go_or_search)
        self.crawl_btn.clicked.connect(self.start_crawl)
        new_tab_btn.clicked.connect(lambda: self.add_new_tab(QUrl("https://www.baidu.com")))

        # 添加到工具栏
        toolbar.addWidget(self.back_btn)
        toolbar.addWidget(self.forward_btn)
        toolbar.addWidget(self.reload_btn)
        toolbar.addWidget(self.home_btn)
        toolbar.addSeparator()
        toolbar.addWidget(self.downloads_btn)
        toolbar.addWidget(self.history_btn)
        toolbar.addWidget(self.bookmarks_btn)
        
        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Minimum)
        toolbar.addWidget(spacer)
        
        toolbar.addWidget(self.url_bar)
        toolbar.addWidget(self.crawl_btn)
        toolbar.addWidget(new_tab_btn)
        parent_layout.addWidget(toolbar)

    def add_new_tab(self, url, title="新标签页"):
        browser = QWebEngineView()
        
        # 创建自定义页面，并传递主窗口引用
        page = CustomWebEnginePage(browser, self)
        browser.setPage(page)
        
        # 设置更真实的用户代理
        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        browser.page().profile().setHttpUserAgent(user_agent)
        
        # 启用所有必要的Web引擎功能
        settings = browser.settings()
        settings.setAttribute(QWebEngineSettings.JavascriptEnabled, True)
        settings.setAttribute(QWebEngineSettings.PluginsEnabled, True)
        settings.setAttribute(QWebEngineSettings.FullScreenSupportEnabled, True)
        settings.setAttribute(QWebEngineSettings.ScrollAnimatorEnabled, True)
        settings.setAttribute(QWebEngineSettings.AutoLoadImages, True)
        settings.setAttribute(QWebEngineSettings.JavascriptCanOpenWindows, True)
        settings.setAttribute(QWebEngineSettings.JavascriptCanAccessClipboard, True)
        settings.setAttribute(QWebEngineSettings.LocalStorageEnabled, True)
        settings.setAttribute(QWebEngineSettings.LocalContentCanAccessRemoteUrls, True)
        settings.setAttribute(QWebEngineSettings.AllowRunningInsecureContent, True)
        settings.setAttribute(QWebEngineSettings.AllowWindowActivationFromJavaScript, True)
        settings.setAttribute(QWebEngineSettings.HyperlinkAuditingEnabled, True)
        settings.setAttribute(QWebEngineSettings.PlaybackRequiresUserGesture, False)

        browser.load(url)
        browser.titleChanged.connect(lambda t: self.update_tab_title(browser, t))
        browser.loadFinished.connect(lambda ok: self.on_load_finished(ok, browser))
        browser.urlChanged.connect(lambda q: self.on_url_changed(browser, q))

        index = self.tab_widget.addTab(browser, title)
        self.tab_widget.setCurrentIndex(index)
        return browser

    def update_tab_title(self, browser, title):
        index = self.tab_widget.indexOf(browser)
        if index != -1:
            truncated = title[:20] + "..." if len(title) > 20 else title
            self.tab_widget.setTabText(index, truncated)

    def on_load_finished(self, success, browser):
        if success:
            self.status_label.setText("页面加载完成")
            # 添加到历史记录
            self.history_manager.add_history(browser.title(), browser.url().toString())
        else:
            self.status_label.setText("加载失败")
        self.update_navigation_buttons()

    def on_url_changed(self, browser, url):
        if browser == self.tab_widget.currentWidget():
            self.url_bar.setText(url.toString())

    def close_tab(self, index):
        if self.tab_widget.count() > 1:
            widget = self.tab_widget.widget(index)
            widget.deleteLater()
            self.tab_widget.removeTab(index)
        else:
            self.tab_widget.clear()
            self.add_new_tab(QUrl("https://www.baidu.com"))

    def on_tab_changed(self, index):
        if index == -1: return
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            self.url_bar.setText(browser.url().toString())
            self.update_navigation_buttons()

    def on_back_clicked(self):
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView): browser.back()

    def on_forward_clicked(self):
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView): browser.forward()

    def on_reload_clicked(self):
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            browser.reload()

    def go_home(self):
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            browser.load(QUrl("https://www.baidu.com"))

    def on_go_or_search(self):
        text = self.url_bar.text().strip()
        if not text: return

        if text.startswith(("http://", "https://")):
            url = QUrl(text)
        elif '.' in text:
            url = QUrl("https://" + text)
        else:
            encoded = requests.utils.quote(text)
            url = QUrl(f"https://www.baidu.com/s?wd={encoded}")

        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView): 
            browser.load(url)

    def start_crawl(self):
        browser = self.tab_widget.currentWidget()
        if not isinstance(browser, QWebEngineView): return

        current_url = browser.url().toString()

        if not self.crawler.can_fetch(current_url):
            reply = QMessageBox.question(
                self, "风险提示",
                f"⚠️ robots.txt 不允许爬取该网站。\n是否继续？",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.No: return

        self.status_label.setText("🔍 正在抓取页面...")
        QTimer.singleShot(100, lambda: self._do_crawl_in_thread(current_url))

    def _do_crawl_in_thread(self, url):
        success, msg = self.crawler.crawl_single_page(url)
        self.status_label.setText(f"{'✅' if success else '❌'} {msg}")
        self.update_data_list()

    def create_data_panel(self):
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(10, 10, 10, 10)

        group = QGroupBox("爬取数据管理")
        group_layout = QVBoxLayout()

        # 数据列表
        self.data_list = QListWidget()
        self.data_list.itemClicked.connect(self.show_data_detail)

        # 数据预览
        self.data_preview = QTextEdit()
        self.data_preview.setReadOnly(True)

        # 操作按钮
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("💾 保存数据")
        clear_btn = QPushButton("🗑️ 清空")
        export_txt = QPushButton("📤 导出训练集")
        export_docx = QPushButton("📄 导出DOCX")
        export_docx.setEnabled(DOCX_AVAILABLE)

        save_btn.clicked.connect(self.save_all_data)
        clear_btn.clicked.connect(self.clear_all_data)
        export_txt.clicked.connect(self.export_training_data)
        export_docx.clicked.connect(self.export_as_docx)

        btn_layout.addWidget(save_btn)
        btn_layout.addWidget(clear_btn)
        btn_layout.addWidget(export_txt)
        btn_layout.addWidget(export_docx)

        group_layout.addWidget(QLabel("已抓取页面"))
        group_layout.addWidget(self.data_list)
        group_layout.addWidget(QLabel("内容预览"))
        group_layout.addWidget(self.data_preview)
        group_layout.addLayout(btn_layout)
        group.setLayout(group_layout)
        layout.addWidget(group)
        return panel

    def update_data_list(self):
        self.data_list.clear()
        for i, d in enumerate(self.crawler.crawled_data):
            item = QListWidgetItem(f"{i+1}. {d['title']} ({d['word_count']}字)")
            item.setData(Qt.UserRole, i)
            self.data_list.addItem(item)

    def show_data_detail(self, item):
        idx = item.data(Qt.UserRole)
        data = self.crawler.crawled_data[idx]
        content = (
            f"📌 标题: {data['title']}\n"
            f"🔗 URL: {data['url']}\n"
            f"📅 时间: {data['timestamp']}\n"
            f"📝 字数: {data['word_count']}, 字符: {data['char_count']}\n"
            f"🔗 内链:{data['internal_links']} 外链:{data['external_links']}\n\n"
            f"{data['full_content']}"
        )
        self.data_preview.setText(content)

    def save_all_data(self):
        path = os.path.join(self.crawler.output_dir, "crawled_data.json")
        try:
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(self.crawler.crawled_data, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, "成功", f"数据已保存到：\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "失败", str(e))

    def clear_all_data(self):
        if QMessageBox.question(self, "确认", "清空所有爬取数据？") == QMessageBox.Yes:
            self.crawler.crawled_data.clear()
            self.data_list.clear()
            self.data_preview.clear()

    def export_training_data(self):
        path = os.path.join(self.crawler.output_dir, "training_data.txt")
        try:
            with open(path, 'w', encoding='utf-8') as f:
                for d in self.crawler.crawled_data:
                    f.write(f"URL: {d['url']}\n")
                    f.write(f"标题: {d['title']}\n")
                    f.write(f"内容:\n{d['full_content']}\n")
                    f.write("\n" + "="*80 + "\n\n")
            QMessageBox.information(self, "成功", f"训练数据已生成：\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "失败", str(e))

    def export_as_docx(self):
        path, _ = QFileDialog.getSaveFileName(self, "保存DOCX", "", "Word文件 (*.docx)")
        if not path: return
        success, msg = self.crawler.export_to_docx(path)
        QMessageBox.information(self, "结果", msg)

    def update_navigation_buttons(self):
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            self.back_btn.setEnabled(browser.history().canGoBack())
            self.forward_btn.setEnabled(browser.history().canGoForward())

    # 新增的浏览器功能
    def new_window(self):
        """新建浏览器窗口"""
        new_browser = ModernBrowser()
        new_browser.show()

    def cut(self):
        """剪切"""
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            browser.triggerPageAction(QWebEnginePage.Cut)

    def copy(self):
        """复制"""
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            browser.triggerPageAction(QWebEnginePage.Copy)

    def paste(self):
        """粘贴"""
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            browser.triggerPageAction(QWebEnginePage.Paste)

    def zoom_in(self):
        """放大页面"""
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            current_zoom = browser.zoomFactor()
            browser.setZoomFactor(min(current_zoom + 0.1, 3.0))

    def zoom_out(self):
        """缩小页面"""
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            current_zoom = browser.zoomFactor()
            browser.setZoomFactor(max(current_zoom - 0.1, 0.25))

    def zoom_reset(self):
        """重置缩放"""
        browser = self.tab_widget.currentWidget()
        if isinstance(browser, QWebEngineView):
            browser.setZoomFactor(1.0)

    def add_bookmark(self):
        """添加书签"""
        self.bookmarks_manager.add_bookmark()

    def show_about(self):
        """显示关于对话框"""
        about_text = """
        <h2>道衍AI浏览器</h2>
        <p><b>版本：</b>2.0 专业版</p>
        <p><b>功能特性：</b></p>
        <ul>
            <li>智能网页爬虫和数据提取</li>
            <li>完整的下载管理系统</li>
            <li>浏览历史记录</li>
            <li>书签管理</li>
            <li>多标签页浏览</li>
            <li>AI训练数据导出</li>
            <li>合规robots.txt检查</li>
        </ul>
        <p><b>技术支持：</b>基于PyQt5和QWebEngine构建</p>
        """
        QMessageBox.about(self, "关于道衍AI浏览器", about_text)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setFont(QFont("Microsoft YaHei", 10))

    # 检查依赖提示
    missing = []
    for pkg, mod in [("requests", "requests"), ("bs4", "bs4")]:
        try:
            __import__(mod)
        except ImportError:
            missing.append(pkg)
    if missing:
        QMessageBox.warning(None, "缺少依赖", f"请安装: pip install {' '.join(missing)}")

    if not SELENIUM_AVAILABLE:
        QMessageBox.information(None, "提示", "Selenium不可用 → 动态页面可能无法抓取")
    if not DOCX_AVAILABLE:
        QMessageBox.information(None, "提示", "DOCX导出功能不可用，请安装 python-docx")

    window = ModernBrowser()
    window.show()
    sys.exit(app.exec_())